import os
import json
from datetime import datetime
from docx import Document

from src.agent_3.llm_client import LLMClient
from src.agent_3.templates.cv_structure import CV_PROMPT_TEMPLATE, CV_JSON_SCHEMA
from src.agent_3.fabrication_guard import FabricationGuard, FabricationError

class CVBuilder:
    def __init__(self, llm_client: LLMClient):
        self.llm_client = llm_client
        self.guard = FabricationGuard()

    async def build_cv(self, row: dict, mapped_skills: dict, master_cv: dict, ats_feedback: str = None) -> dict:
        # Construct Prompt
        prompt = CV_PROMPT_TEMPLATE.format(
            company=row.get("company", ""),
            role_title=row.get("role_title", ""),
            job_url=row.get("job_url", ""),
            required_skills=", ".join(row.get("required_skills", [])),
            job_description=row.get("job_description", "None provided"),
            ats_feedback=ats_feedback or "None",
            master_cv=json.dumps(master_cv, indent=2),
            mapped_accomplishments=json.dumps(mapped_skills, indent=2)
        )
        
        # Call LLM
        generated_cv = await self.llm_client.generate_json(prompt, CV_JSON_SCHEMA)
        
        # Validate Fabrication
        validation = self.guard.validate_cv(generated_cv, mapped_skills, master_cv)
        if not validation["is_valid"]:
            # Retry once with stricter prompt
            strict_prompt = prompt + "\n\nCRITICAL WARNING: Your previous attempt failed fabrication validation. Do NOT invent " + json.dumps(validation["violations"])
            generated_cv = await self.llm_client.generate_json(strict_prompt, CV_JSON_SCHEMA)
            validation = self.guard.validate_cv(generated_cv, mapped_skills, master_cv)
            if not validation["is_valid"]:
                raise FabricationError(f"Fabrication detected: {validation['violations']}")

        # Generate DOCX
        doc_path = self._generate_docx(generated_cv, row.get("company", "Unknown"))
        
        return {
            "cv_doc_link": doc_path,
            "generated_cv": generated_cv,   # ← returned so pipeline can pass to email/DM writers
            "validation_passed": True
        }
        
    def _generate_docx(self, cv_data: dict, company: str) -> str:
        folder_id = os.environ.get('GOOGLE_GENERATED_CVS_FOLDER_ID', 'your_generated_cvs_folder_id')
        
        date_str = datetime.now().strftime("%Y-%m-%d")
        safe_company = "".join([c for c in company if c.isalpha() or c.isdigit() or c==' ']).rstrip()
        filename = f"{safe_company}_APM_CV_{date_str}"
        
        # 1. ALWAYS save a local backup in the project
        local_backup_dir = os.path.join(os.getcwd(), "backups", "cvs")
        os.makedirs(local_backup_dir, exist_ok=True)
        local_filepath = os.path.join(local_backup_dir, filename + ".docx")
        
        doc = Document()
        # Contact Info
        doc.add_heading(cv_data.get("name", ""), 0)
        contact_info = f"{cv_data.get('email', '')} | {cv_data.get('phone', '')} | {cv_data.get('linkedin', '')}"
        doc.add_paragraph(contact_info)
        
        # Summary
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(cv_data.get("summary", ""))
        
        # Experience
        doc.add_heading("Experience", level=1)
        for exp in cv_data.get("experience", []):
            doc.add_heading(f"{exp.get('role')} at {exp.get('company')} ({exp.get('duration')})", level=2)
            for bullet in exp.get("bullets", []):
                doc.add_paragraph(bullet, style='List Bullet')
                
        # Education
        doc.add_heading("Education", level=1)
        for edu in cv_data.get("education", []):
            doc.add_paragraph(f"{edu.get('degree')} - {edu.get('school')} ({edu.get('year')})")
            
        # Skills
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(cv_data.get("skills", [])))
        
        doc.save(local_filepath)
        print(f"INFO: Local backup saved to {local_filepath}")

        # 2. Upload to Google Drive if configured
        if folder_id and folder_id != 'your_generated_cvs_folder_id':
            try:
                from src.shared.google_auth import GoogleDriveAuth
                auth = GoogleDriveAuth()
                drive_service = auth.get_drive_service()
                if drive_service:
                    # 2. Upload the local .docx file directly to bypass service account quota
                    from googleapiclient.http import MediaFileUpload
                    file_metadata = {
                        'name': filename + ".docx",
                        'parents': [folder_id]
                    }
                    media = MediaFileUpload(local_filepath, 
                                            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                                            resumable=True)
                    uploaded_file = drive_service.files().create(body=file_metadata,
                                                               media_body=media,
                                                               fields='id, webViewLink').execute()
                    
                    return uploaded_file.get('webViewLink')
            except Exception as e:
                print(f"Error creating Google Doc: {e}. Returning local path.")
                
        return local_filepath
